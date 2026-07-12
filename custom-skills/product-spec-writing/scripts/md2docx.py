#!/usr/bin/env python3
"""Convert a Markdown product/spec/milestone doc to .docx using python-docx.

Use this when pandoc is unavailable (common on WSL / fresh boxes) and the user
wants a Word deliverable. Install once:  pip install python-docx

Handles: ATX headings (#..######), paragraphs, '-'/'*' bullet lists, and
GitHub-Flavored pipe tables — including cells that use inline
<ul><li>..</li></ul> bullets (the table-cell bullet trick this skill documents).
Inline **bold**/*italic* markers are stripped to plain text for simplicity.

Usage:  python3 md2docx.py INPUT.md [OUTPUT.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def strip_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", text)
    return text.strip()


def cell_lines(cell: str):
    """Turn a table cell into lines, expanding <ul><li>..</li></ul> to bullets."""
    cell = cell.strip()
    if "<li>" in cell:
        items = re.findall(r"<li>(.*?)</li>", cell, flags=re.S)
        return [("bullet", strip_inline(re.sub(r"<.*?>", "", i))) for i in items]
    return [("plain", strip_inline(re.sub(r"<.*?>", "", cell)))]


def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Light Grid Accent 1"
    for r in rows:
        cells = t.add_row().cells
        for i in range(cols):
            raw = r[i] if i < len(r) else ""
            lines = cell_lines(raw)
            cell = cells[i]
            cell.text = ""
            for j, (kind, txt) in enumerate(lines):
                p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
                if kind == "bullet":
                    p.text = "\u2022 " + txt
                else:
                    p.text = txt
    return t


def convert(md_path: str, out_path: str | None = None) -> str:
    md = Path(md_path).read_text(encoding="utf-8")
    out_path = out_path or str(Path(md_path).with_suffix(".docx"))
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # pipe table: a header row followed by a |---|---| separator
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(
            r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]
        ):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = []
            for n, row in enumerate(block):
                if n == 1:  # separator row
                    continue
                cells = [c for c in row.strip().strip("|").split("|")]
                rows.append([c.strip() for c in cells])
            add_table(doc, rows)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            doc.add_heading(strip_inline(m.group(2)), level=min(len(m.group(1)), 4))
        elif re.match(r"^\s*[-*]\s+", line):
            doc.add_paragraph(strip_inline(re.sub(r"^\s*[-*]\s+", "", line)),
                              style="List Bullet")
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph()
        elif line.strip():
            doc.add_paragraph(strip_inline(line))
        i += 1

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("usage: python3 md2docx.py INPUT.md [OUTPUT.docx]")
    saved = convert(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
    print(f"saved {saved}")
