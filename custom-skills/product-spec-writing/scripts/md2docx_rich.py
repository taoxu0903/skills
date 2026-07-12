#!/usr/bin/env python3
"""Rich Markdown -> .docx for formatting-heavy PM docs.

Use INSTEAD of the bundled scripts/md2docx.py when the doc leans on any of:
- **bold** lead-ins / *italic* / `code` inline runs
- <span style="color:rgb(r,g,b)">..</span>  (e.g. Ken's red TBD/NOTE flags)
- numbered step lists (1. 2. ...)
- tab/space-indented sub-bullets
- inline formatting inside table cells

Preserves all of the above: bold/italic/code runs, colored spans as real
colored text, `1.` -> List Number, indented `-` -> List Bullet 2, `>` ->
Intense Quote, GFM pipe tables (bold header row), [label](url) -> blue
underlined label text.

Deliberately injects NO content not present in the source (no synthesized
title / metadata) -- honors Ken's "fully respect the file latest content".

Requires: pip install python-docx
Usage:   cd /tmp && python3 /abs/path/md2docx_rich.py IN.md [OUT.docx]
Run from /tmp with absolute paths to dodge the ~/inspect.py sys.path shadow
gotcha documented in the skill.
"""
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SPAN_RE = re.compile(
    r'<span[^>]*color:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)[^>]*>(.*?)</span>', re.S)
LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
LINK_BLUE = RGBColor(0x0F, 0x62, 0xFE)


def parse_inline(text, style):
    """Return list of (text, style_dict) runs. Recurses for nesting."""
    runs, buf, i, n = [], "", 0, len(text)

    def flush():
        nonlocal buf
        if buf:
            runs.append((buf, dict(style)))
            buf = ""

    while i < n:
        m = SPAN_RE.match(text, i)
        if m:
            flush()
            s2 = dict(style); s2["color"] = (int(m[1]), int(m[2]), int(m[3]))
            runs.extend(parse_inline(m.group(4), s2)); i = m.end(); continue
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j != -1:
                flush()
                s2 = dict(style); s2["code"] = True
                runs.append((text[i + 1:j], s2)); i = j + 1; continue
        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j != -1:
                flush()
                s2 = dict(style); s2["bold"] = True
                runs.extend(parse_inline(text[i + 2:j], s2)); i = j + 2; continue
        if text[i] == "*":
            j = text.find("*", i + 1)
            if j != -1:
                flush()
                s2 = dict(style); s2["italic"] = True
                runs.extend(parse_inline(text[i + 1:j], s2)); i = j + 1; continue
        if text[i] == "[":
            m = LINK_RE.match(text, i)
            if m:
                flush()
                s2 = dict(style); s2["link"] = True
                runs.extend(parse_inline(m.group(1), s2)); i = m.end(); continue
        buf += text[i]; i += 1
    flush()
    return runs


def add_runs(p, runs, force_bold=False):
    if not runs:
        return
    for text, st in runs:
        if text == "":
            continue
        r = p.add_run(text)
        r.bold = bool(st.get("bold")) or force_bold
        r.italic = bool(st.get("italic"))
        if st.get("code"):
            r.font.name = "Consolas"
        if st.get("color"):
            r.font.color.rgb = RGBColor(*st["color"])
        if st.get("link"):
            r.font.color.rgb = LINK_BLUE
            r.underline = True


def add_hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "BFBFBF")):
        bottom.set(qn(k), v)
    pbdr.append(bottom); pPr.append(pbdr)


def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Light Grid Accent 1"
    for ridx, r in enumerate(rows):
        cells = t.add_row().cells
        for c in range(cols):
            raw = r[c] if c < len(r) else ""
            cell = cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, parse_inline(raw.strip(), {}), force_bold=(ridx == 0))
    return t


def convert(md_path, out_path=None):
    md = Path(md_path).read_text(encoding="utf-8")
    out_path = out_path or str(Path(md_path).with_suffix(".docx"))
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # pipe table
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(
                r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = []
            for n, row in enumerate(block):
                if n == 1:
                    continue
                cells = row.strip().strip("|").split("|")
                rows.append([c.strip() for c in cells])
            add_table(doc, rows)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            lvl = max(1, min(len(m.group(1)) - 1, 4))  # ## -> H1, ### -> H2
            h = doc.add_heading("", level=lvl)
            add_runs(h, parse_inline(m.group(2), {}))
            i += 1; continue
        if line.strip() in ("---", "***", "___"):
            add_hrule(doc); i += 1; continue
        mnum = re.match(r"^(\d+)\.\s+(.*)", line)
        if mnum:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, parse_inline(mnum.group(2), {}))
            i += 1; continue
        mbul = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if mbul:
            indent = mbul.group(1).replace("\t", "    ")
            style = "List Bullet 2" if len(indent) >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_runs(p, parse_inline(mbul.group(2), {}))
            i += 1; continue
        if line.lstrip().startswith(">"):
            content = line.lstrip()[1:].strip()
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, parse_inline(content, {}))
            i += 1; continue
        if line.strip():
            p = doc.add_paragraph()
            add_runs(p, parse_inline(line, {}))
        i += 1

    doc.save(out_path)
    return out_path


def verify(out):
    """Structural self-check. Prints counts; asserts no HTML leaked through."""
    assert zipfile.is_zipfile(out), "output is not a valid .docx zip"
    d = Document(out)
    heads = [p.text for p in d.paragraphs
             if p.style.name.startswith("Heading") or p.style.name == "Title"]
    alltext = "\n".join(p.text for p in d.paragraphs)
    for tbl in d.tables:
        for row in tbl.rows:
            for c in row.cells:
                alltext += "\n" + c.text
    leftover = alltext.count("<span") + alltext.count("</span>") + alltext.count("&nbsp;")
    assert leftover == 0, f"HTML leaked into output ({leftover} markers)"
    red = sum(1 for p in d.paragraphs for r in p.runs
              if r.font.color and r.font.color.rgb == RGBColor(255, 0, 0))
    print(f"tables={len(d.tables)} headings={len(heads)} red_runs={red} leftover_html={leftover}")
    print("HEADINGS:", " | ".join(heads))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("usage: python3 md2docx_rich.py INPUT.md [OUTPUT.docx]")
    out = convert(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
    print(f"saved {out}")
    verify(out)
